#!/usr/bin/env python3
import os, re, sys, subprocess, argparse
from docx import Document

LANG = "en"  # will be set from CLI

HEADERS_EN = ["Professional Summary", "Professional Experience", "Education", "Skills", "Competencies", "Languages"]
HEADERS_ES = ["Resumen Profesional", "Experiencia Profesional", "Educación", "Habilidades", "Competencias", "Idiomas"]

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", ".."))
TEMPLATE = os.path.join(ROOT, "cvs", "WordTemplate.docx")
OUTPUT_DIR = os.path.join(ROOT, "cvs", "Generated")

# Template paragraph index map
IDX_MAP = {
    "name": 0,
    "contact": 1,
    "summary_header": 2, "summary": 4,
    "exp_header": 6,
    "job1_title": 8, "job1_company": 9,
    "job1_bullets": list(range(10, 16)),
    "job2_title": 17, "job2_company": 18,
    "job2_bullets": list(range(19, 24)),
    "job3_title": 25, "job3_company": 26,
    "job3_bullets": list(range(27, 36)),
    "edu_header": 37,
    "edu1": 39, "edu2": 40,
    "skills_header": 42,
    "skill_lines": list(range(44, 49)),
    "comp_header": 50,
    "comp_lines": list(range(52, 57)),
    "lang_header": 58, "lang": 60,
}

JOB_KEYS = [
    ("job1_title", "job1_company", "job1_bullets"),
    ("job2_title", "job2_company", "job2_bullets"),
    ("job3_title", "job3_company", "job3_bullets"),
]


def strip_md(text):
    text = re.sub(r'^###\s*', '', text)
    text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
    return text.strip()


def parse_contact(lines):
    """Extract contact info from lines before first ## header."""
    contact = ""
    for line in lines:
        if line.startswith("## "):
            break
        if line.startswith("# ") or line.strip().startswith("---"):
            continue
        s = line.strip()
        if s:
            contact += s + " "
    return contact.strip()


def parse_md(filepath):
    with open(filepath) as f:
        content = f.read()
    lines = content.split("\n")
    sections = {"summary": "", "experience": [], "education": [], "skills": [], "competencies": [], "languages": ""}
    sections["contact"] = parse_contact(lines)
    current = None
    exp_block = None

    for line in lines:
        raw = line
        lower = line.strip().lower().rstrip(":")
        if lower in ("## professional summary", "## summary", "## resumen profesional", "## resumen"):
            current = "summary"; continue
        elif lower in ("## professional experience", "## experience", "## experiencia profesional", "## experiencia"):
            current = "experience"; exp_block = None; continue
        elif lower in ("## education", "## educación"):
            current = "education"; continue
        elif lower in ("## skills", "## habilidades", "## aptitudes"):
            current = "skills"; continue
        elif lower in ("## competencies", "## competencias"):
            current = "competencies"; continue
        elif lower in ("## languages", "## idiomas"):
            current = "languages"; continue

        s = line.strip()
        if not s:
            if current == "experience" and exp_block and exp_block.get("bullets"):
                sections["experience"].append(exp_block)
                exp_block = None
            continue

        if current == "summary":
            sections["summary"] += s + " "
        elif current == "experience":
            if s.startswith("### "):
                if exp_block and exp_block.get("bullets"):
                    sections["experience"].append(exp_block)
                exp_block = {"title": s.lstrip("#").strip(), "company": "", "bullets": []}
            elif exp_block and "|" in s and s.startswith("**"):
                exp_block["company"] = s.replace("**", "").strip()
            elif exp_block is not None and (s.startswith("- ") or s.startswith("* ")):
                exp_block["bullets"].append(s.lstrip("-* ").strip())
        elif current == "education":
            if s == "---":
                continue
            if s.startswith("- ") or s.startswith("* "):
                sections["education"].append(s[2:])
            else:
                sections["education"].append(s)
        elif current == "skills":
            if s.startswith("- ") or s.startswith("* "):
                clean = strip_md(s[2:])
                sections["skills"].append(clean)
        elif current == "competencies":
            if s.startswith("- ") or s.startswith("* "):
                sections["competencies"].append(strip_md(s[2:]))
        elif current == "languages":
            if s.startswith("- ") or s.startswith("* "):
                sections["languages"] += s[2:] + " "
            else:
                sections["languages"] += s + " "

    if exp_block and exp_block.get("bullets"):
        sections["experience"].append(exp_block)
    sections["summary"] = sections["summary"].strip()
    sections["languages"] = strip_md(sections["languages"].strip())
    return sections


def set_run(para, text, run_idx=0, clear_rest=True):
    if not para.runs:
        para.add_run(text)
        return
    if run_idx < len(para.runs):
        para.runs[run_idx].text = text
    if clear_rest:
        for r in para.runs[run_idx + 1:]:
            r.text = ""


def set_skill_line(para, text):
    if ":" in text:
        label, desc = text.split(":", 1)
        if len(para.runs) >= 2:
            para.runs[0].text = label.strip() + ": "
            para.runs[1].text = desc.strip()
            for r in para.runs[2:]:
                r.text = ""
            return
    set_run(para, text)


def set_company_date(para, text):
    if "|" in text:
        parts = [p.strip() for p in text.split("|", 1)]
        if len(para.runs) >= 1:
            para.runs[0].text = parts[0]
        if len(para.runs) >= 3:
            para.runs[1].text = "   |   "
            para.runs[2].text = parts[1] if len(parts) > 1 else ""
            for r in para.runs[3:]:
                r.text = ""
        elif len(para.runs) >= 2:
            para.runs[1].text = "   |   " + (parts[1] if len(parts) > 1 else "")
    else:
        set_run(para, text)


def set_edu_line(para, text):
    text = strip_md(text)
    if "|" in text:
        parts = [p.strip() for p in text.split("|")]
        if len(para.runs) >= 1 and parts:
            para.runs[0].text = parts[0]
        if len(para.runs) >= 2:
            para.runs[1].text = "  ·  " + "  ·  ".join(parts[1:])
            for r in para.runs[2:]:
                r.text = ""
            return
    set_run(para, text)


def set_language_line(para, text):
    if "|" in text:
        parts = [p.strip() for p in text.split("|")]
        pipe_count = text.count("|")
        # Pre-calculate how many runs we need
        new_runs = []
        for part in parts:
            colon_idx = part.find(":")
            if colon_idx >= 0:
                new_runs.append(part[:colon_idx + 1].strip() + " ")
                new_runs.append(part[colon_idx + 1:].strip())
            else:
                new_runs.append(part)
                new_runs.append("")

        # Flatten with separators
        final_texts = []
        for i, run_text in enumerate(new_runs):
            final_texts.append(run_text)
            if i % 2 == 1 and i < len(new_runs) - 1:
                final_texts.append("     |     ")

        for i, r in enumerate(para.runs):
            if i < len(final_texts):
                r.text = final_texts[i]
            else:
                r.text = ""
        return
    set_run(para, text)


def generate_docx(md_path, category):
    print(f"Parsing: {md_path}")
    sections = parse_md(md_path)
    print(f"  Summary: {len(sections['summary'])} chars")
    print(f"  Experience: {len(sections['experience'])} jobs")
    print(f"  Education: {len(sections['education'])} lines")
    print(f"  Skills: {len(sections['skills'])} lines")
    print(f"  Competencies: {len(sections['competencies'])} lines")
    print(f"  Languages: {len(sections['languages'])} chars")

    print(f"Loading template: {TEMPLATE}")
    doc = Document(TEMPLATE)
    paras = doc.paragraphs
    m = IDX_MAP

    # Name
    set_run(paras[m["name"]], "Valeria Páez")

    # Contact info
    if sections.get("contact"):
        set_run(paras[m["contact"]], sections["contact"])
    else:
        set_run(paras[m["contact"]], "")

    # Headers
    headers = HEADERS_ES if LANG == "es" else HEADERS_EN
    for key in ["summary_header", "exp_header", "edu_header", "skills_header", "comp_header", "lang_header"]:
        paras[m[key]].paragraph_format.keep_with_next = True
    set_run(paras[m["summary_header"]], headers[0])
    set_run(paras[m["exp_header"]], headers[1])
    set_run(paras[m["edu_header"]], headers[2])
    set_run(paras[m["skills_header"]], headers[3])
    set_run(paras[m["comp_header"]], headers[4])
    set_run(paras[m["lang_header"]], headers[5])

    # Summary
    if sections["summary"]:
        set_run(paras[m["summary"]], sections["summary"])

    # Experience: replace each job, clear remaining bullet slots
    for job_idx, (title_key, company_key, bullets_key) in enumerate(JOB_KEYS):
        if job_idx < len(sections["experience"]):
            job = sections["experience"][job_idx]
            paras[m[title_key]].paragraph_format.keep_with_next = True
            paras[m[company_key]].paragraph_format.keep_with_next = True
            set_run(paras[m[title_key]], job["title"])
            if job.get("company"):
                set_company_date(paras[m[company_key]], job["company"])
            bullet_idxs = m[bullets_key]
            for bi, bullet_idx in enumerate(bullet_idxs):
                if bi < len(job["bullets"]):
                    set_run(paras[bullet_idx], job["bullets"][bi])
                else:
                    set_run(paras[bullet_idx], "", clear_rest=True)

    # Education
    if len(sections["education"]) > 0:
        paras[m["edu1"]].paragraph_format.keep_with_next = True
        set_edu_line(paras[m["edu1"]], sections["education"][0])
    if len(sections["education"]) > 1:
        paras[m["edu2"]].paragraph_format.keep_with_next = True
        set_edu_line(paras[m["edu2"]], sections["education"][1])

    # Skills
    for si, skill_idx in enumerate(m["skill_lines"]):
        if si < len(sections["skills"]):
            set_skill_line(paras[skill_idx], sections["skills"][si])
        else:
            set_run(paras[skill_idx], "", clear_rest=True)

    # Competencies
    for ci, comp_idx in enumerate(m["comp_lines"]):
        if ci < len(sections["competencies"]):
            set_run(paras[comp_idx], sections["competencies"][ci])
        else:
            set_run(paras[comp_idx], "", clear_rest=True)

    # Languages
    if sections["languages"]:
        set_language_line(paras[m["lang"]], sections["languages"])

    cat_dir = os.path.join(OUTPUT_DIR, category)
    os.makedirs(cat_dir, exist_ok=True)
    safe = category.replace(" ", "_").replace("/", "_")
    suffix = "_es" if LANG == "es" else ""
    out_name = f"CV_Valeria_Paez_Reyes_{safe}{suffix}.docx"
    out_path = os.path.join(cat_dir, out_name)
    doc.save(out_path)
    print(f"✅ DOCX: {out_path}")
    return out_path


def generate_pdf_simple(md_path, out_path):
    """LLM-friendly PDF via md-bookify (minimal metadata)."""
    import shutil
    tmp_md = os.path.join(os.path.dirname(out_path), "__temp_simple.md")
    try:
        shutil.copy2(md_path, tmp_md)
        r = subprocess.run(["md-bookify", tmp_md], capture_output=True, timeout=30)
        if r.returncode == 0:
            pdf_gen = tmp_md.replace(".md", ".pdf")
            if os.path.exists(pdf_gen):
                shutil.move(pdf_gen, out_path)
                print(f"✅ PDF (simple): {out_path}")
                return True
        else:
            print(f"⚠️  md-bookify stderr: {r.stderr.decode()[:200]}")
    except Exception as e:
        print(f"⚠️  md-bookify failed: {e}")
    finally:
        if os.path.exists(tmp_md):
            os.remove(tmp_md)
    return False


def generate_pdf_styled(md_path, out_path):
    """Styled PDF via pandoc + Playwright (Times, page-break control)."""
    tmp_html = os.path.join(os.path.dirname(out_path), "__temp_styled.html")
    try:
        r = subprocess.run(["pandoc", md_path, "-o", tmp_html, "--standalone"],
                           capture_output=True, timeout=30)
        if r.returncode != 0:
            print(f"⚠️  pandoc failed: {r.stderr.decode()[:200]}")
            return False

        with open(tmp_html) as f:
            html = f.read()

        css = """
<style>
@page { margin: 0.6in; }
body {
  font-family: Times, 'Times New Roman', serif;
  font-size: 11.5pt;
  line-height: 1.25;
  color: #1a1a1a;
}
h1 { font-size: 22pt; text-align: center; margin: 0 0 2pt 0; font-weight: bold; }
h1 + p { text-align: center; font-size: 10.5pt; color: #555; margin: 0 0 8pt 0; }
h2 {
  font-size: 11.5pt; margin: 10pt 0 3pt 0;
  text-transform: uppercase; letter-spacing: 1pt;
  border-bottom: 1px solid #333; padding-bottom: 1pt;
  font-weight: bold; page-break-after: avoid;
}
h3 { font-size: 11pt; margin: 6pt 0 0 0; font-weight: bold; page-break-after: avoid; }
h3 + p { font-size: 10.5pt; color: #555; margin: 0 0 2pt 0; }
p { margin: 2pt 0; page-break-inside: avoid; }
ul { margin: 1pt 0 4pt 0; padding-left: 18pt; page-break-inside: avoid; }
li { margin: 0.5pt 0; }
hr { margin: 6pt 0; border: none; border-top: 1px solid #ccc; }
</style>
"""
        html = html.replace("</head>", css + "\n</head>")

        with open(tmp_html, "w") as f:
            f.write(html)

        node_script = f"""
        const {{ chromium }} = require('playwright');
        (async () => {{
          const browser = await chromium.launch();
          const page = await browser.newPage();
          await page.goto('file://{tmp_html}');
          await page.pdf({{ path: '{out_path}', format: 'A4' }});
          await browser.close();
        }})();
        """
        r2 = subprocess.run(["node", "-e", node_script], capture_output=True, timeout=30)
        if r2.returncode != 0:
            print(f"⚠️  Playwright PDF failed: {r2.stderr.decode()[:200]}")
            return False
        print(f"✅ PDF (styled): {out_path}")
        return True
    except Exception as e:
        print(f"⚠️  PDF styled failed: {e}")
        return False
    finally:
        if os.path.exists(tmp_html):
            os.remove(tmp_html)


def generate_pdf(md_path, category):
    cat_dir = os.path.join(OUTPUT_DIR, category)
    os.makedirs(cat_dir, exist_ok=True)
    safe = category.replace(" ", "_").replace("/", "_")
    suffix = "_es" if LANG == "es" else ""
    base = os.path.join(cat_dir, f"CV_Valeria_Paez_Reyes_{safe}{suffix}")
    generate_pdf_simple(md_path, base + "_simple.pdf")
    generate_pdf_styled(md_path, base + "_styled.pdf")


def main():
    global LANG
    parser = argparse.ArgumentParser()
    parser.add_argument("--md", required=True)
    parser.add_argument("--category", required=True)
    parser.add_argument("--lang", default="en", choices=["en", "es"])
    parser.add_argument("--pdf", action="store_true", default=True)
    args = parser.parse_args()
    LANG = args.lang
    if not os.path.exists(args.md):
        print(f"MD not found: {args.md}"); sys.exit(1)
    docx_path = generate_docx(args.md, args.category)
    if args.pdf:
        generate_pdf(args.md, args.category)


if __name__ == "__main__":
    main()
